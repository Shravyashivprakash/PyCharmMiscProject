import os
import re
from docx import Document

from io_utils import clean, norm, extract_label
from config import (
    BASE_INPUT,
    G1_SRS_FILE,
    G1_ICD_FILE,
    G1_SOFTWARE_REQ_DOCX
)

SRS_ID_RE = re.compile(r"SCU_STC_SRS_\d+")


def generate_software_req():
    src = Document(os.path.join(BASE_INPUT, G1_SRS_FILE))
    out = Document()

    table = out.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "ID"
    table.rows[0].cells[1].text = "DESC"

    for tbl in src.tables:
        for row in tbl.rows:
            row_text = " ".join(c.text for c in row.cells)
            m = SRS_ID_RE.search(row_text)
            if not m:
                continue

            rid = m.group()
            r = table.add_row()
            r.cells[0].text = rid
            dst = r.cells[1]

            for cell in row.cells:
                if rid in cell.text:
                    continue

                for p in cell.paragraphs:
                    if p.text.strip():
                        dst.add_paragraph(p.text.strip())

                for nt in cell.tables:
                    t = dst.add_table(rows=len(nt.rows), cols=len(nt.columns))
                    t.style = "Table Grid"
                    for i, rr in enumerate(nt.rows):
                        for j, cc in enumerate(rr.cells):
                            t.rows[i].cells[j].text = cc.text.strip()

    out.save(os.path.join(BASE_INPUT, G1_SOFTWARE_REQ_DOCX))


def extract_srs_arinc():
    doc = Document(os.path.join(BASE_INPUT, G1_SOFTWARE_REQ_DOCX))
    rx, msg = [], []

    for tbl in doc.tables:
        for row in tbl.rows[1:]:
            cell = row.cells[1]
            label = extract_label(cell.text)

            for nt in cell.tables:
                headers = [norm(c.text) for c in nt.rows[0].cells]

                if "receiver" in " ".join(headers):
                    for r in nt.rows[1:]:
                        c = [clean(x.text) for x in r.cells]
                        rx.append({
                            "Label": extract_label(c[1]) or label,
                            "Receiver": c[0],
                            "Interval": c[2]
                        })

                if "bit" in " ".join(headers):
                    for r in nt.rows[1:]:
                        c = [clean(x.text) for x in r.cells]
                        msg.append({
                            "Label": label,
                            "Bit": c[0],
                            "Field": c[1],
                            "Description": c[2],
                            "Definition": c[3]
                        })

    return rx, msg


def extract_icd_arinc():
    doc = Document(os.path.join(BASE_INPUT, G1_ICD_FILE))
    rx, msg = [], []
    current_label = None
    tables = list(doc.tables)
    ti = 0

    for p in doc.paragraphs:
        lbl = extract_label(p.text)
        if lbl:
            current_label = lbl

        if ti < len(tables) and tables[ti]._tbl.getprevious() == p._p:
            t = tables[ti]
            ti += 1
            headers = [norm(c.text) for c in t.rows[0].cells]

            if "receiver" in " ".join(headers):
                for r in t.rows[1:]:
                    c = [clean(x.text) for x in r.cells]
                    rx.append({
                        "Label": extract_label(c[1]),
                        "Receiver": c[0],
                        "Interval": c[2]
                    })

            if "bit" in " ".join(headers):
                for r in t.rows[1:]:
                    c = [clean(x.text) for x in r.cells]
                    msg.append({
                        "Label": current_label,
                        "Bit": c[0],
                        "Field": c[1],
                        "Description": c[2],
                        "Definition": c[3]
                    })

    return rx, msg
