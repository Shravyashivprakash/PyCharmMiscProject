# =========================================================
# G3 IO (USED ONLY BY G3)
# =========================================================

from docx import Document
from openpyxl import load_workbook
import os
from config import REQ_PREFIX


def read_all_text(docx_path):
    """G3 ONLY"""
    doc = Document(docx_path)
    blocks = []

    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    blocks.append(cell.text.strip())

    return blocks


# =========================================================
# G4 IO (USED ONLY BY G4)
# =========================================================

def g4_extract_from_docx(file_path):
    doc = Document(file_path)
    requirements = []

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                rid = row.cells[0].text.strip()
                rtxt = row.cells[1].text.strip()

                if rid.startswith(REQ_PREFIX):
                    requirements.append((rid, rtxt))

    return requirements


def g4_extract_from_excel(file_path):
    wb = load_workbook(file_path, data_only=True)
    requirements = []

    for sheet in wb.worksheets:
        for row in sheet.iter_rows(min_col=1, max_col=2):
            rid, rtxt = row[0].value, row[1].value
            if isinstance(rid, str) and rid.startswith(REQ_PREFIX):
                requirements.append((rid.strip(), str(rtxt).strip()))

    return requirements


def g4_collect_requirements(input_dir):
    """
    Discovers files and extracts requirements.
    No review logic here.
    """
    all_requirements = []

    for fname in os.listdir(input_dir):
        if fname.startswith("~$"):
            continue
        if not ("SRS" in fname.upper() or "SES" in fname.upper() or "TEST" in fname.upper()):
            continue

        path = os.path.join(input_dir, fname)

        if fname.lower().endswith(".docx"):
            all_requirements.extend(g4_extract_from_docx(path))
        elif fname.lower().endswith(".xlsx"):
            all_requirements.extend(g4_extract_from_excel(path))

    return all_requirements
# =========================================================
# G5 IO (USED ONLY BY G5)
# =========================================================

from docx import Document
import re

VALID_ID_PATTERN = re.compile(r"^\s*(SCU_STC_SRS_\d+)\b", re.IGNORECASE)


def extract_requirements(docx_path):
    doc = Document(docx_path)
    requirements = []

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue

            raw_id = row.cells[0].text.strip()
            text = row.cells[1].text.strip()

            if raw_id.lower() in ("requirement id", "srs id", "id"):
                continue

            if not raw_id and text:
                requirements.append({
                    "id": None,
                    "raw_id": "",
                    "text": text
                })
                continue

            if raw_id and not VALID_ID_PATTERN.match(raw_id):
                requirements.append({
                    "id": None,
                    "raw_id": raw_id,
                    "text": text
                })
                continue

            match = VALID_ID_PATTERN.match(raw_id)
            requirements.append({
                "id": match.group(1),
                "raw_id": raw_id,
                "text": text
            })

    return requirements

##### g1 #######
from docx import Document
import re
from config import REQ_ID_RE, LABEL_RE

def extract_docx_text(path):
    doc = Document(path)
    parts = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text.strip():
                    parts.append(c.text)

    return "\n".join(parts)


def normalize_text(text):
    text = text.lower()
    text = re.sub(r"[^\w\s]", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def extract_requirements(text):
    reqs = {}
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    i = 0

    while i < len(lines):
        m = REQ_ID_RE.match(lines[i])
        if m:
            rid = m.group(1)
            buf = []
            i += 1
            while i < len(lines) and not REQ_ID_RE.match(lines[i]):
                buf.append(lines[i])
                i += 1
            reqs[rid] = normalize_text(" ".join(buf))
            continue
        i += 1

    return reqs


def extract_labels(text):
    return {
        m for m in LABEL_RE.findall(text)
        if m.isdigit() and 0 < int(m) <= 377
    }


def clean(text):
    return text.replace("\n", " ").strip() if text else ""


def norm(text):
    return re.sub(r"\s+", " ", str(text).lower().strip()) if text else ""


def parse_bits(text):
    nums = list(map(int, re.findall(r"\d+", str(text))))
    if not nums:
        return None, None
    return min(nums), max(nums)

def extract_label(text):
    """
    Extract single ARINC label from text
    Used by G1.3 ARINC parsing
    """
    if not text:
        return None

    matches = LABEL_RE.findall(text)
    return matches[0] if matches else None
